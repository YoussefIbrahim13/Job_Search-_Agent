"""
backend/parsers/cv_parser.py
============================
CV / résumé file parser.

Supports three formats:
    • PDF   — via PyMuPDF (fitz), using spatial block extraction so multi-column
              layouts are read in correct visual order; falls back to pdfminer
              (linear extraction) if fitz is not installed.
    • DOCX  — via python-docx.
    • TXT   — raw UTF-8 / latin-1 decode.

The module intentionally does NOT call any LLM here. It only extracts raw text,
deterministically harvests known hard skills via local token matching, and builds
a lightweight ``CandidateProfile`` dataclass that the agent layer can then reason
about. Keeping parsing/extraction separate from inference makes unit testing and
future parser swaps trivial, and it takes pressure off the LLM to "guess" skills
that are already verbatim in the document.

Exported symbols
----------------
    CandidateProfile   — dataclass with the parsed information
    parse_cv_bytes     — main entry point; accepts (bytes, filename) → CandidateProfile
"""

from __future__ import annotations

import io
import logging
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Final, List, Optional, Sequence

from backend.core.config import get_settings

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Data model
# ---------------------------------------------------------------------------


@dataclass
class CandidateProfile:
    """
    Structured representation of a parsed CV.

    Fields
    ------
    raw_text        Raw extracted text (truncated to config.max_cv_chars).
    file_name       Original file name (used for logging / error messages).
    detected_title  Best-guess job title inferred from the first non-blank lines.
    word_count      Number of words in raw_text (after truncation).
    skills          Hard skills harvested verbatim from the text via deterministic
                     token matching (not LLM-inferred). Order follows the internal
                     skills dictionary, not order-of-appearance in the document.
    years_experience Total years of professional experience, or None when it
                     could not be determined with reasonable confidence.
    seniority       One of intern/junior/mid/senior/lead/principal, or None when
                     undetermined. None means "unknown" and must render as blank —
                     never substitute a confident-looking default.
    parse_error     Non-empty if parsing produced a recoverable warning.
    """

    raw_text: str = ""
    file_name: str = "unknown"
    detected_title: str = ""
    word_count: int = 0
    skills: List[str] = field(default_factory=list)
    years_experience: Optional[float] = None
    seniority: Optional[str] = None
    parse_error: str = ""

    # Convenience helpers --------------------------------------------------

    def is_empty(self) -> bool:
        """Return True when no meaningful text was extracted."""
        return len(self.raw_text.strip()) < 50

    def summary_for_agent(self) -> str:
        """
        Return a compact, single-string summary suitable for injection into
        the LLM prompt. The text is already truncated; this just wraps it
        with a minimal header so the model understands the context.

        Harvested skills are surfaced explicitly so the LLM treats them as
        ground truth rather than re-deriving them from scratch.
        """
        title_hint = f"Detected title: {self.detected_title}\n\n" if self.detected_title else ""
        skills_hint = (
            f"Verbatim skills detected: {', '.join(self.skills)}\n\n" if self.skills else ""
        )
        return (
            f"=== CV CONTENT ({self.word_count} words) ===\n"
            f"{title_hint}"
            f"{skills_hint}"
            f"{self.raw_text}\n"
            f"=== END CV CONTENT ==="
        )


# ---------------------------------------------------------------------------
# Text normalisation helpers
# ---------------------------------------------------------------------------


def _clean_text(text: str) -> str:
    """Normalise whitespace and remove control characters."""
    # Replace non-breaking spaces and other Unicode whitespace with a regular space
    text = re.sub(r"[\u00a0\u2000-\u200f\u2028\u2029\ufeff]", " ", text)
    # Collapse 3+ consecutive newlines to two (preserve paragraph breaks)
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Strip leading/trailing whitespace from every line
    lines = [line.rstrip() for line in text.splitlines()]
    return "\n".join(lines).strip()


def _truncate(text: str, max_chars: int) -> str:
    """Hard-truncate text and append a marker if content was cut."""
    if len(text) <= max_chars:
        return text
    logger.warning(
        "CV text truncated from %d → %d characters to protect context window.",
        len(text),
        max_chars,
    )
    return text[:max_chars] + "\n\n[... CV TRUNCATED FOR CONTEXT SAFETY ...]"


def _infer_title(text: str) -> str:
    """
    Naively try to extract a job title from the top of the CV.

    Strategy: look at the first 10 non-blank lines; if one of them is short
    (5-60 chars) and comes after a name-like line, use it as the title.
    Returns an empty string when detection fails (the agent will handle it).
    """
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()][:15]
    # Heuristic: short lines that are not email/phone/URL are candidate titles
    url_re = re.compile(r"(http|www\.|@|\+\d)", re.IGNORECASE)
    for line in lines[1:6]:  # skip the very first line (likely the name)
        if 5 <= len(line) <= 70 and not url_re.search(line):
            # Simple title-casing check: most job titles are ≥ 2 words OR
            # contain known keywords
            if re.search(
                r"(engineer|developer|analyst|manager|architect|scientist|"
                r"designer|consultant|lead|intern|specialist|officer|director)",
                line,
                re.IGNORECASE,
            ):
                return line
    return ""


# ---------------------------------------------------------------------------
# Experience extraction
# ---------------------------------------------------------------------------
#
# Design rule for everything below: return None when confidence is low. The
# previous implementation reported the hardcoded string "Professional" for every
# candidate and the UI rendered it as a definitive "LEVEL:" badge. A blank field
# the user can fill in is honest; a confident wrong label is not.
#
# Phase 3 makes these values user-editable; until then they are advisory only.

# "5+ years of experience", "over 3 years' experience", "7 yrs experience"
_EXPLICIT_YEARS_RE = re.compile(
    r"\b(?:over|more\s+than|about|approx(?:imately)?\.?|~)?\s*"
    r"(\d{1,2})\s*\+?\s*(?:years?|yrs?)\b[^.\n]{0,30}?\bexperien\w*",
    re.IGNORECASE,
)

# Date ranges in an experience section: "Jan 2019 - Mar 2022", "2019 – present".
_DATE_RANGE_RE = re.compile(
    r"\b(?:(?P<sm>[A-Za-z]{3,9})\.?\s+)?(?P<sy>(?:19|20)\d{2})\s*"
    r"(?:-|–|—|to|until)\s*"
    r"(?:(?P<em>[A-Za-z]{3,9})\.?\s+)?(?P<ey>(?:19|20)\d{2}|present|current|now)\b",
    re.IGNORECASE,
)

_MONTHS = {
    m: i
    for i, m in enumerate(
        ["jan", "feb", "mar", "apr", "may", "jun",
         "jul", "aug", "sep", "oct", "nov", "dec"],
        start=1,
    )
}

# Ordered most-senior first so the first hit wins.
_SENIORITY_PATTERNS: Sequence[tuple[str, re.Pattern]] = (
    ("principal", re.compile(r"\b(principal|staff|distinguished)\s+\w*\s*engineer\b|\bprincipal\b", re.IGNORECASE)),
    ("lead",      re.compile(r"\b(lead|head\s+of|engineering\s+manager|team\s+lead|tech\s+lead)\b", re.IGNORECASE)),
    ("senior",    re.compile(r"\bsenior\b|\bsr\.?\s", re.IGNORECASE)),
    ("junior",    re.compile(r"\bjunior\b|\bjr\.?\s|\bentry[\s-]level\b|\bgraduate\s+(?:trainee|program)\b", re.IGNORECASE)),
    ("intern",    re.compile(r"\bintern(ship)?\b|\btrainee\b", re.IGNORECASE)),
)


def _extract_years_experience(text: str, current_year: int | None = None) -> Optional[float]:
    """
    Best-effort total years of professional experience.

    Two strategies, in order of reliability:
      1. An explicit self-reported claim ("5+ years of experience").
      2. The span from the earliest to the latest employment date range found.

    Returns None when neither yields a plausible value. Values above 50 are
    treated as a parse artefact rather than a career.
    """
    if not text.strip():
        return None

    explicit = [int(m.group(1)) for m in _EXPLICIT_YEARS_RE.finditer(text)]
    plausible = [y for y in explicit if 0 < y <= 50]
    if plausible:
        return float(max(plausible))

    if current_year is None:
        from datetime import datetime, timezone
        current_year = datetime.now(timezone.utc).year

    starts: List[float] = []
    ends: List[float] = []
    for m in _DATE_RANGE_RE.finditer(text):
        start_year = int(m.group("sy"))
        start_month = _MONTHS.get((m.group("sm") or "")[:3].lower(), 1)

        raw_end = (m.group("ey") or "").lower()
        if raw_end in ("present", "current", "now"):
            end_year, end_month = current_year, 12
        else:
            end_year = int(raw_end)
            end_month = _MONTHS.get((m.group("em") or "")[:3].lower(), 12)

        # Reject impossible or reversed ranges rather than producing a negative.
        if not (1950 <= start_year <= current_year) or end_year < start_year:
            continue
        starts.append(start_year + (start_month - 1) / 12)
        ends.append(end_year + (end_month - 1) / 12)

    if not starts:
        return None

    span = max(ends) - min(starts)
    if span <= 0 or span > 50:
        return None
    return round(span, 1)


def _infer_seniority(text: str, years: Optional[float] = None) -> Optional[str]:
    """
    Best-effort seniority band.

    Prefers an explicit title keyword near the top of the CV (where the current
    role lives); falls back to banding by years of experience. Returns None when
    neither signal is available.
    """
    if not text.strip():
        return None

    # Only the header region — a "senior" mentioned deep in a job description
    # of a past role says nothing about the candidate's current level.
    header = "\n".join(text.splitlines()[:12])
    for label, pattern in _SENIORITY_PATTERNS:
        if pattern.search(header):
            return label

    if years is None:
        return None
    if years < 1:
        return "junior"
    if years < 3:
        return "junior"
    if years < 6:
        return "mid"
    if years < 10:
        return "senior"
    return "lead"


# ---------------------------------------------------------------------------
# Deterministic hard-skills harvesting
# ---------------------------------------------------------------------------
#
# Why this exists: the LLM downstream in the agent graph has a tendency to
# either hallucinate skills the candidate doesn't have, or miss skills that
# are stated explicitly but phrased unusually. Harvesting known tokens here,
# deterministically and locally, removes that guesswork for anything in our
# dictionary and gives the agent verbatim, citable ground truth to reason on
# top of (it can still infer *additional* soft skills/seniority from raw_text).
#
# Matching is intentionally conservative: word-boundary-safe substring
# matching against a curated list, not fuzzy/NLP matching. False negatives
# (an unusual skill we don't recognise) are an acceptable cost; false
# positives (matching "C#" inside "CSS" or "GO" inside "good") are not.

# Skills that contain regex metacharacters (., #, +, /) and therefore need
# the lookaround-based boundary check rather than a plain \b boundary, since
# \b only fires at a transition between a \w and a non-\w character — it does
# not reliably anchor on a token that itself starts or ends with a symbol.
_SKILLS_DICTIONARY: Final[Sequence[str]] = (
    # --- .NET ecosystem (primary focus) ---
    "C#",
    ".NET",
    ".NET Core",
    ".NET Framework",
    "ASP.NET",
    "ASP.NET Core",
    "ASP.NET MVC",
    "Entity Framework",
    "Entity Framework Core",
    "Blazor",
    "Razor",
    "SignalR",
    "Dapper",
    "NHibernate",
    "LINQ",
    "Web API",
    "WPF",
    "WCF",
    "Xamarin",
    "MAUI",
    "NuGet",
    "xUnit",
    "NUnit",
    "MSTest",
    "Moq",
    # --- Databases ---
    "SQL Server",
    "T-SQL",
    "PostgreSQL",
    "MySQL",
    "SQLite",
    "MongoDB",
    "Redis",
    "Elasticsearch",
    "Cassandra",
    "Oracle",
    "DynamoDB",
    "CosmosDB",
    # --- Architecture / design ---
    "Clean Architecture",
    "Modular Monolith",
    "Microservices",
    "Domain-Driven Design",
    "DDD",
    "CQRS",
    "Event Sourcing",
    "SOLID",
    "Design Patterns",
    "REST",
    "RESTful",
    "GraphQL",
    "gRPC",
    "OAuth",
    "OAuth 2.0",
    "JWT",
    "OpenID Connect",
    # --- Other backend / languages ---
    "Java",
    "Spring Boot",
    "Python",
    "Django",
    "Flask",
    "FastAPI",
    "Node.js",
    "Express.js",
    "NestJS",
    "Go",
    "Golang",
    "Rust",
    "C++",
    "PHP",
    "Laravel",
    "Ruby",
    "Ruby on Rails",
    # --- Frontend / web ---
    "JavaScript",
    "TypeScript",
    "React",
    "Angular",
    "Vue.js",
    "HTML",
    "CSS",
    "SASS",
    "Tailwind CSS",
    "Bootstrap",
    "jQuery",
    "Next.js",
    # --- DevOps / cloud / tooling ---
    "Docker",
    "Kubernetes",
    "Azure",
    "AWS",
    "Google Cloud",
    "GCP",
    "Terraform",
    "Ansible",
    "Jenkins",
    "GitHub Actions",
    "GitLab CI",
    "CI/CD",
    "Git",
    "Nginx",
    "RabbitMQ",
    "Kafka",
    "gRPC",
    "Helm",
    # --- AI / data (relevant to Backend AI Engineer roles) ---
    "Machine Learning",
    "Deep Learning",
    "TensorFlow",
    "PyTorch",
    "scikit-learn",
    "LangChain",
    "LangGraph",
    "OpenAI API",
    "Groq",
    "Tavily",
    "NLP",
    "LLM",
    "RAG",
    "Pandas",
    "NumPy",
)


def _build_skill_pattern(skill: str) -> re.Pattern[str]:
    """
    Build a word-boundary-safe regex for a single skill token.

    Plain ``\\b`` boundaries are unreliable for tokens that begin or end with
    a non-word character (``C#``, ``.NET``, ``C++``), because ``\\b`` only
    matches at a transition between a word character and a non-word
    character — a symbol sitting at the edge of the token breaks that
    assumption. Using negative lookaround for "not preceded/followed by an
    alphanumeric character" sidesteps that entirely and works uniformly for
    every token, symbol-edged or not.
    """
    escaped = re.escape(skill)
    return re.compile(rf"(?<![A-Za-z0-9])(?:{escaped})(?![A-Za-z0-9])", re.IGNORECASE)


# Pre-compiled once at import time — the dictionary is static, so there's no
# reason to pay the regex-compilation cost on every CV parsed.
_SKILL_PATTERNS: Final[Sequence[tuple[str, re.Pattern[str]]]] = tuple(
    (skill, _build_skill_pattern(skill)) for skill in _SKILLS_DICTIONARY
)


def _harvest_skills(text: str) -> List[str]:
    """
    Deterministically harvest known hard skills from cleaned CV text.

    Returns skills in dictionary order (not order-of-appearance), de-duplicated,
    using the canonical capitalisation defined in ``_SKILLS_DICTIONARY`` rather
    than whatever casing the candidate happened to use.
    """
    found: List[str] = []
    for canonical_name, pattern in _SKILL_PATTERNS:
        if pattern.search(text):
            found.append(canonical_name)
    return found


# ---------------------------------------------------------------------------
# Format-specific extractors
# ---------------------------------------------------------------------------


def _sort_blocks_reading_order(blocks: List[dict], page_width: float) -> List[dict]:
    """
    Sort text blocks into natural reading order for multi-column layouts.

    Naive PDF text extraction walks the content stream (or a single global
    top-to-bottom sweep) and will happily interleave a left "Experience"
    column with a right "Skills" sidebar whenever their y-coordinates
    overlap — a near-universal trait of templated CVs. That mangles the
    text into something no LLM can reliably reason about.

    The fix: cluster blocks into columns by horizontal position, then read
    each column fully top-to-bottom before moving to the next column,
    left-to-right. Column boundaries are detected from gaps in the
    distribution of block left-edges (``bbox[0]``) relative to page width,
    rather than a fixed pixel threshold, so it scales across page sizes.

    A gap threshold that's too small risks false-splitting a single column
    with indented bullets into two fake "columns"; one that's too large
    risks merging genuinely separate columns. 4% of page width comfortably
    clears typical bullet/list indentation (usually 1-3% of page width)
    while still catching real column gutters (typically 5-10%+).
    """
    if not blocks:
        return blocks

    gap_threshold = page_width * 0.04
    left_edges = sorted(b["bbox"][0] for b in blocks)

    column_starts = [left_edges[0]]
    for previous_x, current_x in zip(left_edges, left_edges[1:]):
        if current_x - previous_x > gap_threshold:
            column_starts.append(current_x)

    def _column_index(x0: float) -> int:
        """Assign a block to the right-most column start it has reached."""
        index = 0
        for candidate_index, start_x in enumerate(column_starts):
            if x0 >= start_x - 1.0:  # small tolerance for sub-pixel jitter
                index = candidate_index
        return index

    for block in blocks:
        block["_column"] = _column_index(block["bbox"][0])

    # Column left-to-right, then top-to-bottom within a column, then left-to-right
    # as a final tiebreaker for blocks that happen to share an exact y-position.
    blocks.sort(key=lambda b: (b["_column"], round(b["bbox"][1], 1), b["bbox"][0]))
    return blocks


def _extract_pdf_blocks_fitz(data: bytes) -> str:
    """
    Extract text from a PDF using PyMuPDF, reading each page via spatial
    blocks sorted into column-aware reading order (see
    ``_sort_blocks_reading_order``) rather than relying on raw stream order.
    """
    import fitz  # PyMuPDF

    doc = fitz.open(stream=data, filetype="pdf")
    try:
        page_texts: List[str] = []
        for page in doc:
            page_dict = page.get_text("dict")
            text_blocks = []
            for block in page_dict.get("blocks", []):
                # type == 0 is a text block; type == 1 is an image block.
                # We only want text — images contribute nothing extractable here.
                if block.get("type") != 0:
                    continue
                block_text = "".join(
                    span["text"]
                    for line in block.get("lines", [])
                    for span in line.get("spans", [])
                ).strip()
                if block_text:
                    text_blocks.append({"bbox": block["bbox"], "text": block_text})

            ordered_blocks = _sort_blocks_reading_order(text_blocks, page.rect.width)
            page_texts.append("\n".join(b["text"] for b in ordered_blocks))

        return "\n\n".join(page_texts)
    finally:
        doc.close()


def _extract_pdf_pdfminer(data: bytes) -> str:
    """
    Fallback PDF extraction via pdfminer.six.

    pdfminer's high-level API does linear, stream-order extraction and has
    no equivalent spatial-block control surface, so multi-column layouts may
    still mangle here. This path only runs when PyMuPDF isn't installed, so
    it trades layout fidelity for availability — acceptable as a fallback,
    not as the primary path.
    """
    from pdfminer.high_level import extract_text_to_fp
    from pdfminer.layout import LAParams

    output = io.StringIO()
    extract_text_to_fp(
        io.BytesIO(data),
        output,
        laparams=LAParams(),
        output_type="text",
        codec="utf-8",
    )
    return output.getvalue()


def _extract_pdf(data: bytes) -> str:
    """Extract text from a PDF, preferring PyMuPDF's spatial-block pipeline."""
    try:
        return _extract_pdf_blocks_fitz(data)
    except ImportError:
        logger.debug("PyMuPDF not available, falling back to pdfminer (linear extraction).")

    try:
        return _extract_pdf_pdfminer(data)
    except ImportError:
        raise RuntimeError(
            "No PDF library found. Install either 'PyMuPDF' or 'pdfminer.six':\n"
            "    pip install PyMuPDF\n"
            "  or\n"
            "    pip install pdfminer.six"
        )


def _extract_docx(data: bytes) -> str:
    """Extract text from a .docx file using python-docx."""
    try:
        from docx import Document  # python-docx
    except ImportError:
        raise RuntimeError(
            "python-docx is required for .docx parsing.\n"
            "Install it with:  pip install python-docx"
        )

    doc = Document(io.BytesIO(data))
    paragraphs = [para.text for para in doc.paragraphs]
    # Also grab text from tables (common in modern CVs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.append(cell.text)
    return "\n".join(paragraphs)


def _extract_txt(data: bytes) -> str:
    """Decode plain-text bytes, trying UTF-8 then latin-1."""
    try:
        return data.decode("utf-8")
    except UnicodeDecodeError:
        return data.decode("latin-1", errors="replace")


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


def parse_cv_bytes(file_bytes: bytes, file_name: str) -> CandidateProfile:
    """
    Parse a CV from raw bytes.

    Parameters
    ----------
    file_bytes  Raw binary content of the uploaded file.
    file_name   Original file name including extension (e.g. ``"john_doe_cv.pdf"``).

    Returns
    -------
    CandidateProfile
        Always returns a profile object; ``parse_error`` is non-empty on soft
        failures (the agent can still attempt to work with partial text).

    Raises
    ------
    ValueError   When the file extension is not supported.
    RuntimeError When a required parsing library is missing.
    """
    settings = get_settings()
    suffix = Path(file_name).suffix.lower()

    # PRIVACY: CV filenames routinely embed the candidate's real name
    # ("Jane_Doe_CV.pdf"), so log the type and size at INFO and keep the name
    # itself at DEBUG. Security-rejection paths in routes.py intentionally still
    # log the filename — there the provenance is the point.
    logger.info("Parsing CV file: type=%s (%d bytes)", suffix or "(none)", len(file_bytes))
    logger.debug("CV filename: %r", file_name)

    parse_error = ""
    raw_text = ""

    try:
        if suffix == ".pdf":
            raw_text = _extract_pdf(file_bytes)
        elif suffix in (".docx", ".doc"):
            raw_text = _extract_docx(file_bytes)
        elif suffix in (".txt", ".text", ".md"):
            raw_text = _extract_txt(file_bytes)
        else:
            raise ValueError(
                f"Unsupported file type '{suffix}'. "
                "Please upload a PDF, DOCX, or TXT file."
            )
    except (ValueError, RuntimeError):
        raise
    except Exception as exc:  # noqa: BLE001
        # Soft failure: record the error but return a partial profile so the
        # agent layer can surface a useful error message to the user.
        parse_error = f"Parsing error: {exc}"
        logger.exception("Unexpected error while parsing '%s'.", file_name)

    # Normalise and truncate -------------------------------------------------
    raw_text = _clean_text(raw_text)
    raw_text = _truncate(raw_text, settings.max_cv_chars)
    word_count = len(raw_text.split())
    detected_title = _infer_title(raw_text)
    skills = _harvest_skills(raw_text)
    years_experience = _extract_years_experience(raw_text)
    seniority = _infer_seniority(raw_text, years_experience)

    profile = CandidateProfile(
        raw_text=raw_text,
        file_name=file_name,
        detected_title=detected_title,
        word_count=word_count,
        skills=skills,
        years_experience=years_experience,
        seniority=seniority,
        parse_error=parse_error,
    )

    if profile.is_empty() and not parse_error:
        profile.parse_error = "CV appears to be empty or contained no extractable text."
        logger.warning("Empty CV after parsing '%s'.", file_name)
    else:
        # PRIVACY: never log CV-derived content (detected title, skill values,
        # raw text) above DEBUG. The UI tells the user their CV is "processed
        # locally, never stored" — writing their job title and skill list into
        # stdout logs quietly breaks that promise. Counts only at INFO.
        logger.info(
            "CV parsed successfully: %d words, title hint %s, %d skills detected.",
            word_count,
            "present" if detected_title else "absent",
            len(skills),
        )
        logger.debug(
            "CV parse detail: title=%r skills=%s",
            detected_title,
            skills,
        )

    return profile